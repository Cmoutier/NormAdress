"""Injection des champs de fusion Word dans le document client."""
from __future__ import annotations
import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CHAMPS_FUSION = ["Formule", "L1", "L2", "L3", "L4", "L5", "L6", "ID_CLIENT"]

# Marqueurs que l'opérateur peut placer dans son courrier Word
# NormAdress les remplace par les champs de fusion réels
MARQUEURS = {
    "{{Formule}}": "Formule",
    "{{L1}}": "L1",
    "{{L2}}": "L2",
    "{{L3}}": "L3",
    "{{L4}}": "L4",
    "{{L5}}": "L5",
    "{{L6}}": "L6",
    "{{ID_CLIENT}}": "ID_CLIENT",
    # Alias courants
    "{{adresse}}": "L1",
    "{{destinataire}}": "L1",
}


def _creer_mergefield(nom: str) -> OxmlElement:
    """Crée un élément XML MERGEFIELD Word."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f' MERGEFIELD "{nom}" ')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = f"«{nom}»"
    r.append(t)
    fld.append(r)
    return fld


def _remplacer_marqueur_dans_run(run, marqueur: str, champ: str):
    """Remplace le texte d'un run par un MERGEFIELD."""
    if marqueur not in run.text:
        return False
    before, _, after = run.text.partition(marqueur)
    run.text = before
    # Insérer le MERGEFIELD après ce run
    fld = _creer_mergefield(champ)
    run._r.addnext(fld)
    if after:
        r_after = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = after
        r_after.append(t)
        fld.addnext(r_after)
    return True


def _zone_adresse_par_marqueurs(doc: Document) -> bool:
    """
    Cherche les marqueurs {{Ln}} dans le texte et les remplace.
    Retourne True si au moins un marqueur a été trouvé.
    """
    found = False
    for para in doc.paragraphs:
        for run in para.runs:
            for marqueur, champ in MARQUEURS.items():
                if marqueur in run.text:
                    _remplacer_marqueur_dans_run(run, marqueur, champ)
                    found = True
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for marqueur, champ in MARQUEURS.items():
                            if marqueur in run.text:
                                _remplacer_marqueur_dans_run(run, marqueur, champ)
                                found = True
    return found


def _injecter_bloc_adresse_en_fin(doc: Document):
    """
    Si aucun marqueur n'est trouvé, ajoute un bloc adresse en fin de document
    avec tous les champs de fusion.
    """
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("--- Bloc adresse NormAdress ---").bold = True

    for champ in ["L1", "L2", "L3", "L4", "L5", "L6"]:
        p = doc.add_paragraph()
        run = p.add_run()
        fld = _creer_mergefield(champ)
        run._r.addnext(fld)

    # Formule sur paragraphe séparé
    p_formule = doc.add_paragraph()
    run_f = p_formule.add_run()
    fld_f = _creer_mergefield("Formule")
    run_f._r.addnext(fld_f)


def injecter_champs_fusion(docx_bytes: bytes) -> bytes:
    """
    Prend le fichier Word client en bytes, injecte les champs de fusion,
    retourne le fichier modifié en bytes.
    """
    buf_in = io.BytesIO(docx_bytes)
    doc = Document(buf_in)

    found = _zone_adresse_par_marqueurs(doc)
    if not found:
        _injecter_bloc_adresse_en_fin(doc)

    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out.read()
